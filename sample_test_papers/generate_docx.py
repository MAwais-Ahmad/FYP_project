import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_from_txt(txt_path, docx_path):
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        if line_str.startswith("===") or line_str.startswith("---"):
            continue
            
        if line_str.startswith("EXAM PAPER:") or line_str.startswith("SECTION") or line_str.startswith("PART"):
            p = doc.add_paragraph()
            run = p.add_run(line_str)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(31, 78, 121)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line_str.startswith("Q") or line_str.startswith("1.") or line_str.startswith("2.") or line_str.startswith("3.") or line_str.startswith("4.") or line_str.startswith("5.") or line_str.startswith("6.") or line_str.startswith("7.") or line_str.startswith("8."):
            p = doc.add_paragraph()
            run = p.add_run(line_str)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 32, 96)
        elif line_str.startswith("Correct Answer:") or line_str.startswith("Key Points:"):
            p = doc.add_paragraph()
            run = p.add_run(line_str)
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(38, 128, 0)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line_str)
            run.font.size = Pt(10.5)
            
    doc.save(docx_path)
    print(f"Saved docx to {docx_path}")

base_dir = r"d:\Projects\FYP_project\sample_test_papers"
create_docx_from_txt(os.path.join(base_dir, "sample_paper_1_ai_cognitive_systems.txt"), os.path.join(base_dir, "sample_paper_1_ai_cognitive_systems.docx"))
create_docx_from_txt(os.path.join(base_dir, "sample_paper_2_information_security.txt"), os.path.join(base_dir, "sample_paper_2_information_security.docx"))
