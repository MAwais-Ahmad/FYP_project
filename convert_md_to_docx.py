import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def convert_md_to_docx(md_filepath, docx_filepath):
    doc = Document()

    # Set normal margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    def flush_table(t_lines):
        if not t_lines:
            return
        # Parse table markdown
        rows = []
        for line in t_lines:
            line = line.strip()
            if not line.startswith('|'):
                continue
            parts = [p.strip() for p in line.split('|')[1:-1]]
            # Check if delimiter row
            if all(re.match(r'^:?-+:?$', p) for p in parts if p):
                continue
            rows.append(parts)

        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    # Clean inline bold/code formatting
                    clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', val)
                    clean_text = re.sub(r'`(.*?)`', r'\1', clean_text)
                    cell.text = clean_text

                    # Format header row
                    if r_idx == 0:
                        set_cell_background(cell, '1F4E79') # Navy Blue
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.size = Pt(10)
                    else:
                        if r_idx % 2 == 1:
                            set_cell_background(cell, 'F2F4F8') # Light zebra stripe
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9.5)

        doc.add_paragraph() # Spacing

    for line in lines:
        line_str = line.strip()

        # Table detection
        if line_str.startswith('|'):
            in_table = True
            table_lines.append(line_str)
            continue
        elif in_table:
            in_table = False
            flush_table(table_lines)
            table_lines = []

        if not line_str:
            continue

        # Headings
        if line_str.startswith('# '):
            h = doc.add_heading(level=1)
            run = h.add_run(line_str[2:].strip())
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
        elif line_str.startswith('## '):
            h = doc.add_heading(level=2)
            run = h.add_run(line_str[3:].strip())
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(46, 117, 182)
        elif line_str.startswith('### '):
            h = doc.add_heading(level=3)
            run = h.add_run(line_str[4:].strip())
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(89, 89, 89)
        elif line_str.startswith('#### '):
            h = doc.add_heading(level=4)
            run = h.add_run(line_str[5:].strip())
            run.font.size = Pt(11)
            run.font.bold = True
        elif line_str.startswith('* ') or line_str.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            content = line_str[2:].strip()
            # Render bold sections inside bullets
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)
        elif re.match(r'^\d+\.\s', line_str):
            p = doc.add_paragraph(style='List Number')
            content = re.sub(r'^\d+\.\s', '', line_str).strip()
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)
        elif line_str == '---':
            continue
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line_str)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)

    if in_table:
        flush_table(table_lines)

    doc.save(docx_filepath)
    print(f"Successfully generated {docx_filepath}")

if __name__ == '__main__':
    src = 'word_doc/AITA_FYP_Final_Report_Document.md' if os.path.exists('word_doc/AITA_FYP_Final_Report_Document.md') else 'AITA_FYP_Final_Report_Document.md'
    out = 'word_doc/AITA_FYP_Final_Report_Document.docx'
    convert_md_to_docx(src, out)
